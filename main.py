from docx import Document

document = Document()

document.add_heading('Заголовок', 1)
document.add_paragraph('А это обычный текст')
document.add_heading('А вот это - маркированный список', 2)
document.add_paragraph('Пункт 1', style='List Bullet')
document.add_paragraph('Пункт 2', style='List Bullet')
document.add_heading('А вот это - нумерованный список', 2)
document.add_paragraph('Пункт 1', style='List Number')
document.add_paragraph('Пункт 2', style='List Number')

document.save('test.docx')
